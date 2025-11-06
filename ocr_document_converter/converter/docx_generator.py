"""
Word Document Generator Module
Creates Word documents from OCR results
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import logging


class DocxGenerator:
    """Generate Word documents from OCR text"""
    
    def __init__(self):
        """Initialize document generator"""
        pass
    
    def create_simple_document(self, text, output_path):
        """
        Create a simple Word document with plain text
        
        Args:
            text: Text content
            output_path: Path to save .docx file
            
        Returns:
            str: Path to created document
        """
        try:
            doc = Document()
            
            # Add text content
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Save document
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logging.info(f"Created simple document: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logging.error(f"Error creating simple document: {e}")
            raise
    
    def create_formatted_document(self, ocr_results, output_path, preserve_colors=True):
        """
        Create a Word document with formatting from OCR results
        
        Args:
            ocr_results: Dictionary with OCR data including text and word boxes
            output_path: Path to save .docx file
            preserve_colors: Whether to apply detected colors
            
        Returns:
            str: Path to created document
        """
        try:
            doc = Document()
            
            # Get text and word data
            text = ocr_results.get('text', '')
            words = ocr_results.get('words', [])
            
            if not text.strip():
                # Empty document
                doc.add_paragraph("(No text detected)")
            elif not words or not preserve_colors:
                # Simple text without color formatting
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        para = doc.add_paragraph()
                        run = para.add_run(line)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    else:
                        doc.add_paragraph()
            else:
                # Format with colors (Phase 2)
                self._create_colored_document(doc, words, ocr_results)
            
            # Save document
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logging.info(f"Created formatted document: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logging.error(f"Error creating formatted document: {e}")
            raise
    
    def _create_colored_document(self, doc, words, ocr_results):
        """
        Create document with colored text based on detected colors
        
        Args:
            doc: Document object
            words: List of word dictionaries with text, position, and color
            ocr_results: Full OCR results
        """
        try:
            # Group words by line based on y-coordinate
            lines = self._group_words_into_lines(words)
            
            for line_words in lines:
                if not line_words:
                    doc.add_paragraph()  # Empty line
                    continue
                
                para = doc.add_paragraph()
                
                for word_data in line_words:
                    text = word_data.get('text', '')
                    color = word_data.get('color', (0, 0, 0))
                    font_size = self._estimate_font_size(word_data.get('height', 12))
                    
                    # Add word with formatting
                    run = para.add_run(text + ' ')  # Add space between words
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(font_size)
                    
                    # Apply color
                    if color and color != (0, 0, 0):  # Non-black color
                        from docx.shared import RGBColor
                        r, g, b = color
                        run.font.color.rgb = RGBColor(r, g, b)
            
            logging.info(f"Created document with {len(lines)} lines and color formatting")
            
        except Exception as e:
            logging.error(f"Error creating colored document: {e}")
            # Fallback to simple text
            text = ocr_results.get('text', '')
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
    
    def _group_words_into_lines(self, words, line_threshold=10):
        """
        Group words into lines based on vertical position
        
        Args:
            words: List of word dictionaries with 'y' coordinate
            line_threshold: Maximum y-difference to consider same line
            
        Returns:
            list: List of lists, each containing words in a line
        """
        if not words:
            return []
        
        # Sort words by y-coordinate (top to bottom)
        sorted_words = sorted(words, key=lambda w: w.get('y', 0))
        
        lines = []
        current_line = [sorted_words[0]]
        current_y = sorted_words[0].get('y', 0)
        
        for word in sorted_words[1:]:
            word_y = word.get('y', 0)
            
            if abs(word_y - current_y) <= line_threshold:
                # Same line
                current_line.append(word)
            else:
                # New line
                # Sort current line by x-coordinate (left to right)
                current_line.sort(key=lambda w: w.get('x', 0))
                lines.append(current_line)
                
                current_line = [word]
                current_y = word_y
        
        # Add last line
        if current_line:
            current_line.sort(key=lambda w: w.get('x', 0))
            lines.append(current_line)
        
        return lines
    
    def _estimate_font_size(self, height):
        """
        Estimate font size in points from pixel height
        
        Args:
            height: Character height in pixels
            
        Returns:
            int: Font size in points
        """
        # Rough conversion: 1 point ≈ 1.33 pixels at 96 DPI
        # Clamp between 8 and 20 points for readability
        size = max(8, min(20, int(height * 0.75)))
        return size
    
    def create_multi_page_document(self, page_results, output_path, preserve_colors=True):
        """
        Create a Word document from multiple pages of OCR results
        
        Args:
            page_results: List of OCR result dictionaries (one per page)
            output_path: Path to save .docx file
            preserve_colors: Whether to apply detected colors
            
        Returns:
            str: Path to created document
        """
        try:
            doc = Document()
            
            for page_num, ocr_result in enumerate(page_results, 1):
                text = ocr_result.get('text', '')
                words = ocr_result.get('words', [])
                
                if not text.strip():
                    doc.add_paragraph(f"(Page {page_num} - No text detected)")
                elif not words or not preserve_colors:
                    # Simple text without colors
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            para = doc.add_paragraph()
                            run = para.add_run(line)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                        else:
                            doc.add_paragraph()
                else:
                    # Format with colors
                    self._create_colored_document(doc, words, ocr_result)
                
                # Add page break if not last page
                if page_num < len(page_results):
                    doc.add_page_break()
            
            # Save document
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logging.info(f"Created multi-page document with {len(page_results)} pages: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logging.error(f"Error creating multi-page document: {e}")
            raise
    
    def add_colored_text(self, paragraph, text, color_rgb):
        """
        Add text with color to a paragraph
        
        Args:
            paragraph: docx Paragraph object
            text: Text to add
            color_rgb: Tuple of (R, G, B) values (0-255)
        """
        run = paragraph.add_run(text)
        if color_rgb:
            r, g, b = color_rgb
            run.font.color.rgb = RGBColor(r, g, b)
        return run
    
    def add_formatted_text(self, paragraph, text, font_size=None, bold=False, italic=False, color_rgb=None):
        """
        Add formatted text to a paragraph
        
        Args:
            paragraph: docx Paragraph object
            text: Text to add
            font_size: Font size in points
            bold: Whether text is bold
            italic: Whether text is italic
            color_rgb: Tuple of (R, G, B) values (0-255)
        """
        run = paragraph.add_run(text)
        
        if font_size:
            run.font.size = Pt(font_size)
        
        if bold:
            run.font.bold = True
        
        if italic:
            run.font.italic = True
        
        if color_rgb:
            r, g, b = color_rgb
            run.font.color.rgb = RGBColor(r, g, b)
        
        return run
    
    def add_image_to_document(self, doc, image_path, width_inches=None):
        """
        Add an image to the document
        
        Args:
            doc: Document object
            image_path: Path to image file
            width_inches: Optional width in inches
        """
        try:
            if width_inches:
                doc.add_picture(str(image_path), width=Inches(width_inches))
            else:
                doc.add_picture(str(image_path))
            
            logging.info(f"Added image to document: {Path(image_path).name}")
            
        except Exception as e:
            logging.error(f"Error adding image to document: {e}")
